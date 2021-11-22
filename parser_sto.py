import os
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


class DocxWriter:
    """Класс, который записывает данные в Docx документ"""
    def __init__(self):
        self.doc = docx.Document()
        self.picture_count = 0
        normal_style = self.doc.styles['Normal']

        font = normal_style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        paragraph_format = normal_style.paragraph_format
        paragraph_format.first_line_indent = Cm(1.25)
        paragraph_format.space_after = 0

        self.keep_with_next = False


    def write_to_doc(self, text):
        if text == '\*':
            self.keep_with_next = not self.keep_with_next
            return

        if text.startswith('\\!'):
            self.add_bold_text_paragraph(text[2:], self.keep_with_next)
        elif text.strip().startswith('&') and text.strip().endswith('&'):
            self.add_image(text[1:-1], self.keep_with_next)
        else:
            self.text_paragraph(self.keep_with_next).add_run(text)
        #elif text.startswith('')


    def text_paragraph(self, keep_with_next=False):
        par = self.doc.add_paragraph()
        par.paragraph_format.keep_with_next = keep_with_next
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        return par


    def add_bold_text_paragraph(self, text, keep_with_next=False):
        par = self.text_paragraph(keep_with_next)
        par.add_run(text).bold = True


    def add_image(self, text='', keep_with_next=False):
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.first_line_indent = 0
        par.paragraph_format.keep_with_next = keep_with_next

        run = par.add_run()
        self.picture_count += 1
        picture_path = f'Screenshots/{self.picture_count}.png'

        if '_MEIPASS2' in os.environ:
            filename = os.path.join(os.environ['_MEIPASS2'], picture_path)

        run.add_picture(picture_path)
        image_text = par.add_run(f'\nРисунок {self.picture_count}' +
                                (f' – {text}' if len(text) else ''))
        font = image_text.font
        font.size = Pt(12)


    def save_docx(self, name):
        self.doc.save(name)
