import itertools
import re
from pathlib import Path
from typing import Any, Callable

import docx
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.section import Section
from docx.shared import Cm, Mm, Pt
from docx.table import Table
from docx.text.font import Font
from docx.text.paragraph import Paragraph
from docx.text.parfmt import ParagraphFormat
from docx.text.run import Run

from .constants import (
    SCREENSHOTS_DIRECTORY_NAME,
    TEMPLATE_CONSTANTS,
    TEMPLATE_NAME,
)
from .exceptions import MissingImage, MissingImageName
from .utils import PictureMeta, get_config_setting, get_project_directory_name


class DocxWriter:
    """Class to write data to `docx` document."""

    def __init__(self) -> None:
        """Init document."""
        self.doc: Document = self.parse_template()
        self.picture_count = 0
        self.keep_with_next = False

        normal_style = self.doc.styles["Normal"]
        self.init_font(normal_style.font)
        self.init_format(normal_style.paragraph_format)
        self.init_a4(self.doc.sections[0])

    @property
    def document_picture_name(self) -> str:
        """Standard name of picture."""
        return f"Рисунок {self.picture_count}"

    @property
    def format_mapping(self) -> dict[str, Callable]:
        """Return format mapping.

        Each format symbol has a corresponding format function. So, this
        mapping just defines which format method will be used according to
        format symbol found on the line.

        """
        return {
            "\\*": self.switch_keep_with_next,
            "\\!": self.add_bold_run,
            "&": self.add_image,
            "": self.add_run,
        }

    @property
    def image_extensions(self) -> itertools.chain:
        """List of all image extensions."""
        image_extensions = [
            ".png",
            ".jpeg",
            ".jpg",
            ".gif",
        ]

        return itertools.chain(
            image_extensions,
            map(str.upper, image_extensions),
        )

    def parse_template(self) -> Document:
        """Return parsed template list."""
        template = docx.Document(
            f"{get_project_directory_name()}/"
            f"{get_config_setting(TEMPLATE_NAME, section='template')}.docx"
        )

        for paragraph in template.paragraphs:
            self.format_template_paragraph(paragraph)

        for table in template.tables:
            self.format_template_table(table)

        return template

    def format_template_paragraph(self, paragraph: Paragraph) -> None:
        """Return formatted paragraph using settings for template page."""
        text = paragraph.text

        if "{" not in text or "}" not in text:
            return

        paragraph.text = text.format(**{
            constant: get_config_setting(constant, section="template")
            for constant in TEMPLATE_CONSTANTS
        })

    def format_template_table(self, table: Table) -> None:
        """Return formatted cells of table using settings for template page."""
        for row_num in range(len(table.rows)):
            for col_num in range(len(table.columns)):
                cell = table.cell(row_idx=row_num, col_idx=col_num)
                self.format_template_paragraph(cell.paragraphs[0])

    def init_a4(self, section: Section) -> None:
        """Init A4 document."""
        section.page_height = Mm(297)
        section.page_width = Mm(210)

    def init_font(self, font: Font) -> None:
        """Init font of the document."""
        font.name = "Times New Roman"
        font.size = Pt(14)

    def init_format(self, paragraph_format: ParagraphFormat) -> None:
        """Init format of the document."""
        paragraph_format.first_line_indent = Cm(1.25)
        paragraph_format.space_after = 0

    def write_to_doc(self, text: str) -> None:
        """Write text to document."""
        format_symbol = self.parse_format_symbols(text)
        format_method = self.format_mapping[format_symbol]
        clean_text = text.strip(format_symbol).strip()
        format_method(clean_text)

    def get_text_paragraph(
        self,
        align_center: bool = False,
    ) -> Paragraph:
        """Return prepared text paragraph."""
        par = self.doc.add_paragraph()
        par.paragraph_format.keep_with_next = self.keep_with_next

        if align_center:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        return par

    def add_run(self, text: str) -> Run:
        """Add paragraph with plain text."""
        par = self.get_text_paragraph()
        run = par.add_run(text)
        return run

    def add_bold_run(self, text: str) -> None:
        """Add paragraph with bold text."""
        run = self.add_run(text)
        run.bold = True

    def add_image(self, text: str = "") -> None:
        """Add image to paragraph."""
        par = self.get_text_paragraph(align_center=True)
        par.paragraph_format.first_line_indent = 0
        run = par.add_run()

        self.picture_count += 1
        meta = self.get_picture_meta(text)
        path = self.get_picture_path(meta.name)
        run.add_picture(path)

        description = f" – {meta.description}" if len(meta.description) else ""
        image_text = par.add_run(
            f"\n{self.document_picture_name}"
            f"{description}"
        )

        font = image_text.font
        font.size = Pt(12)

    def switch_keep_with_next(self, *args: Any) -> None:
        """Switch `self.keep_with_next`."""
        self.keep_with_next = not self.keep_with_next

    def parse_format_symbols(self, text: str) -> str:
        """Return format symbol from given text.

        If format symbols are not presented in the string, then return just
        empty string.

        """
        for symbol in self.format_mapping:
            if text.startswith(symbol):
                return symbol
        return ""

    def get_picture_meta(self, text: str) -> PictureMeta:
        """Return parsed name and description of the picture.

        Raises:
            MissingImageName: If image name was not provided.

        """
        found_match = re.match(r"<(.+?)>(.*)", text)

        if found_match is None:
            raise MissingImageName(image=text)

        return PictureMeta(
            name=found_match.group(1),
            description=found_match.group(2),
        )

    def get_picture_path(self, picture_name: str) -> str:
        """Return path to picture.

        Raises:
            MissingImage: If there are no images with this name.

        """
        base_directory = Path(get_project_directory_name())
        pictures_folder = get_config_setting(SCREENSHOTS_DIRECTORY_NAME)

        path = base_directory / pictures_folder / picture_name
        extension_provided = "." in picture_name

        if extension_provided and not path.is_file():
            raise MissingImage(
                imagename=picture_name,
                picturesfolder=pictures_folder,
            )
        elif not extension_provided:
            path = self.get_path_with_extension(path)

        return str(path)

    def get_path_with_extension(self, path_wo_extension: Path) -> Path:
        """Find extension of the given file and return full path.

        Raises:
            MissingImage: If no images were found with this name.

        """
        for extension in self.image_extensions:
            if (path := path_wo_extension.with_suffix(extension)).is_file():
                return path

        raise MissingImage(
            imagename=path_wo_extension.name,
            picturesfolder=get_config_setting(SCREENSHOTS_DIRECTORY_NAME),
        )

    def save_docx(self, name: str) -> None:
        """Save `docx` file."""
        self.doc.save(name)
